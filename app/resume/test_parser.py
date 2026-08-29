import tempfile
import unittest
from pathlib import Path

from docx import Document

from .parser import extract_resume_text


class ResumeParserTests(unittest.TestCase):
    def test_docx_text_is_extracted(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "resume.docx"
            document = Document()
            document.add_heading("Backend Developer")
            document.add_paragraph("Python, FastAPI, PostgreSQL")
            document.save(path)
            text = extract_resume_text(str(path))
        self.assertIn("Backend Developer", text)
        self.assertIn("FastAPI", text)

    def test_unsupported_format_is_rejected(self) -> None:
        with self.assertRaises(ValueError):
            extract_resume_text("resume.txt")


if __name__ == "__main__":
    unittest.main(verbosity=2)
